from docx import Document
import PyPDF2

''' this in the CvToSimpleText class working explanation
    in this class, we define methods to extract text from different document formats
    we use the python-docx library to handle .docx files
    we use the PyPDF2 library to handle .pdf files
'''

class CvToSimpleText():
    '''
        this in the extractTextFromDocx function working explanation
        first we load the document and extract the text from it
        then we return the extracted text as a string
    '''
    def extractTextFromDocx(filePath : str|None) -> str:
        doc = Document(filePath)
        textParts = []
        seen = set()
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text and text not in seen:
                textParts.append(text)
                seen.add(text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text and text not in seen:
                        textParts.append(text)
                        seen.add(text)
        return "\n".join(textParts)

    '''
    this in the extractTextFromPdf function working explanation
    first we load the PDF document and extract the text from it
    then we return the extracted text as a string
    '''
    def extractTextFromPdf(filePath : str|None) -> str:
        textParts = []
        seen = set()

        with open(filePath, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                pageText = page.extract_text()
                if pageText:
                    for line in pageText.splitlines():
                        cleanLine = line.strip()
                        if cleanLine and cleanLine not in seen:
                            textParts.append(cleanLine)
                            seen.add(cleanLine)
        return "\n".join(textParts)


''' remove this in Production! '''
# if __name__ == "__main__":
#     # Example using pdf 
#     pdf_text = CvToSimpleText.extractTextFromPdf(r"demoData\srijanraycv.pdf")
#     print(pdf_text)
    
#     # Example using docx
#     docx_text = CvToSimpleText.extractTextFromDocx(r"demoData\srijanraycv.docx")
#     print(docx_text)